from io import BytesIO
from typing import Optional
from docx import Document
from PyPDF2 import PdfReader
from bs4 import BeautifulSoup
from striprtf.striprtf import rtf_to_text
import logging

logger = logging.getLogger(__name__)

class DocumentParser:
    """Parser for different document types"""
    
    @staticmethod
    def parse_document(file_content: bytes, mime_type: str) -> Optional[str]:
        """
        Parse document content based on mime type
        
        :param file_content: Binary content of the file
        :param mime_type: MIME type of the document
        :return: Extracted text or None if parsing failed
        """
        try:
            if mime_type == 'application/pdf':
                return DocumentParser._parse_pdf(file_content)
            elif mime_type in ['application/msword', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document']:
                return DocumentParser._parse_docx(file_content)
            elif mime_type == 'text/html':
                return DocumentParser._parse_html(file_content)
            elif mime_type == 'application/rtf':
                return DocumentParser._parse_rtf(file_content)
            else:
                logger.warning(f"Unsupported mime type: {mime_type}")
                return None
        except Exception as e:
            logger.error(f"Error parsing document: {e}", exc_info=True)
            return None

    @staticmethod
    def _parse_pdf(content: bytes) -> str:
        """Extract text from PDF file"""
        with BytesIO(content) as pdf_file:
            pdf = PdfReader(pdf_file)
            text = []
            for page in pdf.pages:
                text.append(page.extract_text())
            return "\n".join(text)

    @staticmethod
    def _parse_docx(content: bytes) -> str:
        """Extract text from DOCX file"""
        try:
            with BytesIO(content) as docx_file:
                doc = Document(docx_file)

                paragraphs = [paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()]

                tables_text = []
                for table in doc.tables:
                    for row in table.rows:
                        row_text = ' '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
                        if row_text:
                            tables_text.append(row_text)
                
                all_text = paragraphs + tables_text
                result = '\n'.join(all_text)
                
                return result
        except Exception as e:
            logger.error(f"Error parsing DOCX: {e}", exc_info=True)
            raise

    @staticmethod
    def _parse_html(content: bytes) -> str:
        """Extract text from HTML file"""
        html_text = content.decode('utf-8', errors='ignore')
        soup = BeautifulSoup(html_text, 'lxml')
        for script in soup(["script", "style"]):
            script.decompose()
        text = soup.get_text()
        lines = (line.strip() for line in text.splitlines())
        chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
        return "\n".join(chunk for chunk in chunks if chunk)

    @staticmethod
    def _parse_rtf(content: bytes) -> str:
        """Extract text from RTF file"""
        try:
            rtf_text = content.decode('utf-8', errors='ignore')
            plain_text = rtf_to_text(rtf_text)
            return plain_text
        except Exception as e:
            logger.error(f"Error parsing RTF: {e}", exc_info=True)
            raise 